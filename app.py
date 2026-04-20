import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import io
from groq import Groq

# Configuración de la IA
client = Groq(api_key=st.secrets["GROQ_API_KEY"])

def extraer_datos_estacion_trabajo(file):
    res = {"paciente": {}, "mediciones": {}, "doppler_clinico": []}
    try:
        xls = pd.ExcelFile(file)
        
        # 1. HOJA ECODATO (Datos fijos y números)
        df_eco = pd.read_excel(xls, "Ecodato", header=None).astype(str)
        res["paciente"]["nombre"] = df_eco.iloc[0, 1].replace("nan", "").strip().upper()
        res["paciente"]["fecha"] = df_eco.iloc[1, 1].replace("nan", "").split(" ")[0]
        
        # Buscamos S/C en la columna E (Dubois u otros)
        res["paciente"]["sc"] = df_eco.iloc[10, 4].replace("nan", "").strip()

        # Mapeo de mediciones principales
        for r in range(5, 25):
            sigla = str(df_eco.iloc[r, 0]).strip().upper()
            valor = df_eco.iloc[r, 1].replace("nan", "").strip()
            if sigla and valor and sigla != "NAN":
                res["mediciones"][sigla] = valor

        # 2. HOJA DOPPLER (Lógica de Marcas 'X')
        if "Doppler" in xls.sheet_names:
            df_dop = pd.read_excel(xls, "Doppler")
            # El doctor pone 'x' en las columnas de 'Sí' o 'No' (Insuficiencia)
            for _, row in df_dop.iterrows():
                valvula = str(row.iloc[0]).strip()
                velocidad = str(row.iloc[1]).replace("nan", "").strip()
                
                # Verificamos si marcó 'x' en Insuficiencia (columna 4 o 5 del Excel)
                si_insuf = str(row.iloc[4]).lower().strip() == 'x'
                
                if valvula != "nan":
                    estado = "CON INSUFICIENCIA" if si_insuf else "SIN INSUFICIENCIA"
                    vel_txt = f"VEL: {velocidad} CM/S" if velocidad else ""
                    res["doppler_clinico"].append(f"{valvula.upper()}: {estado} {vel_txt}")
                    
    except Exception as e:
        st.error(f"Error en lectura de Excel: {e}")
    return res

def redaccion_medica_senior(datos):
    # Prompt ultra-estricto para evitar lenguaje coloquial
    prompt = f"""
    TRANSCRIPTOR MÉDICO CARDIÓLOGO.
    DATOS: {datos['mediciones']}
    DOPPLER: {datos['doppler_clinico']}
    
    TAREA: REDACTA LOS 'HALLAZGOS' Y LA 'CONCLUSIÓN'.
    - TODO EN MAYÚSCULAS.
    - NO USAR 'SUGIERE', 'RECOMIENDA' NI 'ADJUNTO'.
    - LENGUAJE SECO, DESCRIPTIVO Y TÉCNICO.
    - SI DDVI > 56 MM ESCRIBE: 'DILATACIÓN DEL VENTRÍCULO IZQUIERDO'.
    - SI FA < 28 % ESCRIBE: 'DETERIORO DE LA FUNCIÓN SISTÓLICA'.
    """
    completion = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "system", "content": "Eres un sistema de dictado médico. No hablas, solo transcribes técnica."},
                  {"role": "user", "content": prompt}],
        temperature=0
    )
    return completion.choices[0].message.content

def generar_documento_final(datos, texto_ia, pdf_imgs):
    doc = Document()
    
    # 1. Cabecera Profesional
    title = doc.add_heading('Ecocardiograma 2D y Doppler Cardíaco Color', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Cuadro de Paciente
    p_info = doc.add_paragraph()
    run = p_info.add_run(f"PACIENTE: {datos['paciente']['nombre']}\n")
    run.bold = True
    run.font.size = Pt(12)
    p_info.add_run(f"FECHA: {datos['paciente']['fecha']} | S/C: {datos['paciente']['sc']} m²")

    # 3. Cuerpo del Informe (IA)
    partes = texto_ia.upper().split("CONCLUSIÓN")
    
    doc.add_heading('HALLAZGOS', level=1)
    h_text = doc.add_paragraph(partes[0].replace("HALLAZGOS:", "").strip())
    h_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if len(partes) > 1:
        doc.add_heading('CONCLUSIÓN', level=1)
        c_text = doc.add_paragraph(partes[1].replace(":", "").strip())
        c_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        c_text.bold = True

    # 4. Grilla de Imágenes 4x2
    if pdf_imgs:
        doc.add_page_break()
        doc.add_heading('ANEXO DE IMÁGENES', level=1)
        try:
            pdf_imgs.seek(0)
            pdf = fitz.open(stream=pdf_imgs.read(), filetype="pdf")
            imgs = []
            for page in pdf:
                for img_info in page.get_images():
                    imgs.append(io.BytesIO(pdf.extract_image(img_info[0])["image"]))
            
            table = doc.add_table(rows=0, cols=2)
            for i in range(0, min(len(imgs), 8), 2): # Máximo 8 imágenes (4 filas x 2 col)
                row_cells = table.add_row().cells
                for j in range(2):
                    if i + j < len(imgs):
                        para = row_cells[j].paragraphs[0]
                        para.add_run().add_picture(imgs[i+j], width=Inches(2.8))
        except: pass

    # 5. FIRMA BLINDADA (Tabla invisible al final)
    doc.add_paragraph("\n\n")
    firma_table = doc.add_table(rows=1, cols=2)
    firma_table.columns[0].width = Inches(4.5)
    celda = firma_table.rows[0].cells[1]
    f_p = celda.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_p.add_run("__________________________\n").bold = True
    f_p.add_run("FIRMA Y SELLO DEL MÉDICO").bold = True

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# --- UI STREAMLIT ---
st.set_page_config(page_title="CardioReport Senior", page_icon="🩺")
st.title("Estación de Trabajo Cardíaca 🩺")

xl = st.file_uploader("1. Subir Excel (Ecodato + Doppler)", type="xlsx")
pd_f = st.file_uploader("2. Subir PDF de Imágenes", type="pdf")

if xl and pd_f:
    if st.button("GENERAR INFORME MÉDICO FINAL"):
        with st.spinner("Procesando datos y marcas del médico..."):
            info = extraer_datos_completos = extraer_datos_estacion_trabajo(xl)
            relato = redaccion_medica_senior(info)
            archivo = generar_documento_final(info, relato, pd_f)
            
            st.success("Informe generado correctamente.")
            st.download_button("📥 DESCARGAR INFORME (.DOCX)", archivo, 
                             file_name=f"Informe_{info['paciente']['nombre']}.docx")
